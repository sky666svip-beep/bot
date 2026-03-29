import os
import uuid
import logging
import fitz
import json
import re
from flask import Blueprint, request, jsonify, current_app
from docx import Document
from sqlalchemy import func, or_
from werkzeug.exceptions import HTTPException
from werkzeug.utils import secure_filename
from app.models import UserHistory, QuestionBank, Poetry, PoetryAnalysis, Formula, Vocabulary, Idiom
from app.extensions import db
from app.services.llm_service import extract_text_from_image, analyze_essay, generate_study_plan, generate_exam_questions, generate_poetry_analysis
from app.services.async_task import task_mgr
from flask_login import login_required, current_user

api_bp = Blueprint('api', __name__)

def _get_owner():
    """统一的 owner 标识：登录用户用 ID，未登录用 session cookie"""
    return str(current_user.id) if current_user.is_authenticated else request.cookies.get('session', 'anon')

UPLOAD_FOLDER = os.path.join(os.getcwd(), 'instance', 'uploads')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# 安全配置：允许的文件后缀
ALLOWED_IMAGE_EXTS = {'png', 'jpg', 'jpeg', 'webp', 'bmp', 'gif','tiff','tif'}
ALLOWED_DOC_EXTS = {'pdf', 'docx', 'txt'}

def allowed_file(filename, allowed_set):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_set

@api_bp.errorhandler(Exception)
def handle_global_error(e):
    """统一拦截 API 异常，防止对外抛出 500 HTML页面"""
    logging.error(f"API Error: {str(e)}", exc_info=True)
    code = e.code if isinstance(e, HTTPException) else 500
    return jsonify({"success": False, "message": "服务器内部错误" if code == 500 else str(e)}), code

# === 通用任务轮询端点 ===
@api_bp.route('/task/<task_id>/status', methods=['GET'])
def poll_task(task_id):
    """查询异步任务状态，毫秒级响应，不占线程资源"""
    return jsonify(task_mgr.get_status(task_id, owner=_get_owner()))

# 历史记录路由
@api_bp.route('/history', methods=['GET'])
@login_required
def get_history():
    histories = UserHistory.query.filter_by(user_id=current_user.id)\
        .order_by(UserHistory.created_at.desc()).limit(10).all()
    return jsonify([h.to_dict() for h in histories])

@api_bp.route('/history-data', methods=['GET'])
@login_required
def get_history_data():
    filter_mistake = request.args.get('filter') == 'mistake'
    query = UserHistory.query.filter_by(user_id=current_user.id)
    if filter_mistake:
        query = query.filter_by(is_mistake=True)
    
    limit = 200 if filter_mistake else 50
    histories = query.order_by(UserHistory.created_at.desc()).limit(limit).all()
    return jsonify([h.to_dict() for h in histories])

@api_bp.route('/history/<int:hid>/toggle', methods=['POST'])
@login_required
def toggle_history_status(hid):
     #切换历史记录的错题(mistake)状态,添加or删除错题
    record = UserHistory.query.get_or_404(hid)
    
    # 权限校验：只能修改自己的记录
    if record.user_id != current_user.id:
        return jsonify({"success": False, "message": "无权操作"}), 403
        
    record.is_mistake = not record.is_mistake
    db.session.commit()
    return jsonify({"success": True, "new_status": record.is_mistake})

# 文档上传路由
@api_bp.route('/upload-doc', methods=['POST'])
def upload_document():
    file = request.files.get('file')
    if not file or not file.filename:
        return jsonify({"success": False, "message": "无效文件"})
    if not allowed_file(file.filename, ALLOWED_DOC_EXTS):
        return jsonify({"success": False, "message": "不支持的文件格式"}), 400

    # 1. 使用 secure_filename 清洗文件名，防止路径穿越漏洞
    safe_filename = secure_filename(file.filename) or "unnamed_document.tmp"
    
    temp_path = os.path.join(UPLOAD_FOLDER, f"temp_{uuid.uuid4()}_{safe_filename}")
    file.save(temp_path)
    ext = safe_filename.lower().split('.')[-1]

    def _parse_doc():
        """后台线程：解析文档"""
        try:
            text = ""
            if ext == 'pdf':
                with fitz.open(temp_path) as doc:
                    text = "".join([page.get_text() for page in doc])
            elif ext == 'docx':
                doc = Document(temp_path)
                text = "\n".join([p.text for p in doc.paragraphs])
                for table in doc.tables:
                    for row in table.rows:
                        text += "\n" + " ".join([cell.text for cell in row.cells])
            elif ext == 'txt':
                with open(temp_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            else:
                return {"success": False, "message": "不支持的格式"}
            if not text.strip():
                return {"success": False, "message": "文档内容为空"}
            return {"success": True, "full_text": text}
        except Exception as e:
            return {"success": False, "message": f"解析错误: {str(e)}"}
        finally:
            if os.path.exists(temp_path): os.remove(temp_path)

    task_id = task_mgr.submit(_parse_doc, owner=_get_owner())
    return jsonify({"success": True, "task_id": task_id}), 202

@api_bp.route('/dashboard', methods=['GET'])
@login_required
def get_dashboard_data():
    # 1. 热力图
    daily = db.session.query(func.date(UserHistory.created_at), func.count(UserHistory.id))\
        .filter(UserHistory.user_id == current_user.id)\
        .group_by(func.date(UserHistory.created_at)).all()
    heatmap = [[str(d), c] for d, c in daily]
    
    # 2. 饼图SELECT category, COUNT(*) FROM user_history GROUP BY category
    cats = db.session.query(UserHistory.category, func.count(UserHistory.id))\
        .filter(UserHistory.user_id == current_user.id)\
        .group_by(UserHistory.category).all()
    pie = [{"name": c, "value": n} for c, n in cats]
    
    return jsonify({"heatmap": heatmap, "pie": pie})

# === 作文批改API ===

@api_bp.route('/essay/correct', methods=['POST'])
def correct_essay_api():
    data = request.json or {}
    text = data.get('text', '')
    if len(text) < 5:
        return jsonify({"success": False, "message": "内容太短"})
    essay_type = data.get('type', 'chinese')

    def _correct():
        result = analyze_essay(text, essay_type)
        return {"success": True, "data": result}

    task_id = task_mgr.submit(_correct, owner=_get_owner())
    return jsonify({"success": True, "task_id": task_id}), 202

# === 纯文字识别 API (供作文拍照批改使用) ===
@api_bp.route('/ocr-image', methods=['POST'])
def ocr_image_api():
    file = request.files.get('file')
    if not file: return jsonify({'success': False, 'message': '无文件'})
    if not allowed_file(file.filename, ALLOWED_IMAGE_EXTS):
        return jsonify({'success': False, 'message': '不支持的图片格式'})

    temp_path = os.path.join(UPLOAD_FOLDER, f"ocr_{uuid.uuid4()}.jpg")
    file.save(temp_path)

    def _ocr():
        try:
            return {'success': True, 'text': extract_text_from_image(temp_path)}
        finally:
            if os.path.exists(temp_path): os.remove(temp_path)

    task_id = task_mgr.submit(_ocr, owner=_get_owner())
    return jsonify({'success': True, 'task_id': task_id}), 202

# === 学习计划板块 ===

@api_bp.route('/study-plan/generate', methods=['POST'])
def generate_plan_api():
    data = request.json or {}
    if 'grade' not in data or 'duration' not in data:
        return jsonify({"success": False, "message": "缺少必要参数"})

    def _gen_plan():
        return {"success": True, "data": generate_study_plan(data)}

    task_id = task_mgr.submit(_gen_plan, owner=_get_owner())
    return jsonify({"success": True, "task_id": task_id}), 202

@api_bp.route('/study-plan/weakness-analysis', methods=['GET'])
@login_required
def analyze_weakness_api():
    """智能分析错题本：统计最近50条错题，找出高频学科/标签"""
    mistakes = UserHistory.query.filter_by(is_mistake=True, user_id=current_user.id)\
        .order_by(UserHistory.created_at.desc()).limit(50).all()
        
    if not mistakes:
        return jsonify({"success": True, "weakness": "暂无错题数据"})
        
    # 简化统计逻辑
    counts = {}
    for m in mistakes:
        c = m.category or "其他"
        counts[c] = counts.get(c, 0) + 1
        
    top = sorted(counts, key=counts.get, reverse=True)[:3]
    return jsonify({"success": True, "weakness": f"建议重点突破：{', '.join(top)}"})

# === 模拟考试模块 ===
@api_bp.route('/simulation/generate', methods=['POST'])
def generate_simulation_api():
    """生成试题接口"""
    criteria = request.json

    def _gen_exam():
        return {"success": True, "questions": generate_exam_questions(criteria)}

    task_id = task_mgr.submit(_gen_exam, owner=_get_owner())
    return jsonify({"success": True, "task_id": task_id}), 202

@api_bp.route('/simulation/submit', methods=['POST'])
@login_required
def submit_simulation_api():
    """提交并保存考试结果 (异步批量版)"""
    data = request.json or {}
    results = data.get('results', [])
    if not results:
        return jsonify({"success": True, "saved_ids": []})

    uid = current_user.id
    app = current_app._get_current_object()

    def _submit_exam():
        from app.services.nlp_service import nlp_engine

        saved_ids = []
        history_records = []
        new_questions = []

        texts_to_encode = [nlp_engine.clean_prefix(item.get('question', '')) for item in results]
        std_texts = [nlp_engine.standardize_text(t) for t in texts_to_encode]

        if nlp_engine.model:
            embeddings = nlp_engine.model.encode(std_texts, batch_size=32, convert_to_numpy=True, normalize_embeddings=True)
        else:
            embeddings = [[] for _ in results]

        from app.services.answer_engine import extract_core_numbers, _parse_options
        for i, item in enumerate(results):
            history = UserHistory(
                question=item.get('question'), answer=item.get('answer'), reason=item.get('reason'),
                source='模拟考试', category=item.get('category', '其他'), is_mistake=False, user_id=uid
            )
            history_records.append(history)

            # --- 二阶精密去重校验 (与 answer_engine.py 保持同频) ---
            existings = QuestionBank.query.filter_by(std_q=std_texts[i]).all()
            is_dup = False
            if existings:
                user_nums = extract_core_numbers(texts_to_encode[i])
                user_opts_str = json.dumps(item.get('options'), ensure_ascii=False) if item.get('options') else None
                for ex in existings:
                    if user_nums != extract_core_numbers(ex.question):
                        continue
                    ex_opts_str = json.dumps(_parse_options(ex.options), ensure_ascii=False) if ex.options else None
                    if user_opts_str != ex_opts_str:
                        continue
                    is_dup = True
                    print(f"⚠️ [跳过入库] 模拟考试题目已存在 ID: {ex.id}")
                    break
            
            if not is_dup:
                opts = item.get('options')
                new_q = QuestionBank(
                    question=texts_to_encode[i], std_q=std_texts[i], answer=item.get('answer'),
                    reason=item.get('reason'), options=json.dumps(opts, ensure_ascii=False) if opts else None,
                    category=item.get('category', '模拟考试'), embedding=json.dumps(embeddings[i].tolist()) if len(embeddings[i])>0 else None
                )
                new_questions.append(new_q)

        db.session.add_all(history_records)
        db.session.flush()
        for i, history in enumerate(history_records):
            saved_ids.append({"temp_id": results[i].get('temp_id'), "db_id": history.id})

        db.session.bulk_save_objects(new_questions)
        db.session.commit()

        for q in new_questions:
            if q.embedding:
                nlp_engine.add_to_index(q.question, json.loads(q.embedding), q.answer, q.reason, q.options)

        return {"success": True, "saved_ids": saved_ids}

    task_id = task_mgr.submit(_submit_exam, app=app, owner=str(uid))
    return jsonify({"success": True, "task_id": task_id}), 202

# === 古诗词鉴赏模块 ===

@api_bp.route('/poetry/search', methods=['POST'])
def search_poetry():
    keyword = request.json.get('keyword', '').strip()
    if not keyword:
        return jsonify({"success": False, "message": "关键词为空"})

    # 1. 查库 (Exact or Fuzzy 模糊匹配 Title、Author 或 Content)
    poetry = Poetry.query.filter(or_(Poetry.title == keyword, Poetry.author == keyword)).first()
    if not poetry:
        poetry = Poetry.query.filter(or_(Poetry.title.like(f"%{keyword}%"), Poetry.author.like(f"%{keyword}%"), Poetry.content.like(f"%{keyword}%"))).first()
    # 2. 如果库里有诗且有赏析 → 同步返回（毫秒级，无需异步）
    if poetry:
        analysis = PoetryAnalysis.query.filter_by(poetry_id=poetry.id).first()
        if analysis:
            return jsonify({
                "success": True,
                "source": "database",
                "data": {
                    "title": poetry.title,
                    "author": f"{poetry.dynasty or ''} {poetry.author}",
                    "content": poetry.content,
                    "translation": analysis.translation,
                    "appreciation": analysis.appreciation,
                    "annotations": analysis.to_dict()['annotations']
                }
            })

    # 3. 需要 LLM 生成 → 提交到后台线程
    app = current_app._get_current_object()
    poetry_id = poetry.id if poetry else None
    poetry_title = poetry.title if poetry else None
    poetry_author = poetry.author if poetry else None

    def _gen_poetry():
        gen = generate_poetry_analysis(keyword)
        if not isinstance(gen, dict) or 'error' in gen or not gen.get('title') or not gen.get('content'):
            return {"success": False, "message": "未找到相关诗词或解析失败"}

        # 入库
        _poetry_id = poetry_id
        _poetry_title = poetry_title
        _poetry_author = poetry_author
        if not _poetry_id:
            raw_auth = gen.get('author', '佚名')
            dynasty, author = "", raw_auth
            m = re.match(r'^\[(.*?)\]\s*(.*)', raw_auth)
            if m: dynasty, author = m.groups()

            new_poetry = Poetry(title=gen.get('title', keyword), author=author, dynasty=dynasty, content=gen.get('content', ''))
            db.session.add(new_poetry)
            db.session.flush()
            _poetry_id = new_poetry.id
            _poetry_title = new_poetry.title
            _poetry_author = new_poetry.author

        new_analysis = PoetryAnalysis(
            poetry_id=_poetry_id,
            translation=gen.get('translation', ''),
            appreciation=gen.get('appreciation', ''),
            annotations=json.dumps(gen.get('annotations', []), ensure_ascii=False),
            title=_poetry_title,
            author=_poetry_author
        )
        db.session.add(new_analysis)
        db.session.commit()
        return {"success": True, "source": "llm", "data": gen}

    task_id = task_mgr.submit(_gen_poetry, app=app, owner=_get_owner())
    return jsonify({"success": True, "task_id": task_id}), 202

@api_bp.route('/poetry/suggest', methods=['GET'])
def suggest_poetry():
    """输入联想：按关键词模糊匹配作者、标题或诗句内容，返回作品列表"""
    keyword = request.args.get('q', '').strip()
    if not keyword or len(keyword) < 1:
        return jsonify({"data": []})

    results = Poetry.query.filter(
        or_(Poetry.title.like(f"%{keyword}%"), Poetry.author.like(f"%{keyword}%"), Poetry.content.like(f"%{keyword}%"))
    ).limit(30).all()

    return jsonify({
        "data": [
            {"id": p.id, "title": p.title, "author": f"{p.dynasty or ''} {p.author}"}
            for p in results
        ]
    })

# === 公式大全模块 ===

@api_bp.route('/formulas', methods=['GET'])
def get_formulas():
    """获取公式列表（支持分页、筛选）"""
    # 筛选参数
    grade = request.args.get('grade', '')
    category = request.args.get('category', '')
    keyword = request.args.get('keyword', '')
    page = request.args.get('page', 1, type=int)
    per_page = request.args.get('per_page', 20, type=int)
    
    query = Formula.query
    
    # 学段筛选
    if grade:
        query = query.filter(Formula.grade.like(f"%{grade}%"))
    # 学科筛选
    if category:
        query = query.filter(Formula.category == category)
    # 关键词搜索（名称模糊匹配）
    if keyword:
        query = query.filter(or_(
            Formula.name.like(f"%{keyword}%"),
            Formula.tags.like(f"%{keyword}%")
        ))
    
    # 分页
    pagination = query.order_by(Formula.id).paginate(page=page, per_page=per_page, error_out=False)
    
    # 获取所有学段和学科用于筛选器
    all_grades = db.session.query(Formula.grade).distinct().all()
    all_categories = db.session.query(Formula.category).distinct().all()
    
    return jsonify({
        "success": True,
        "data": [f.to_dict() for f in pagination.items],
        "total": pagination.total,
        "page": page,
        "pages": pagination.pages,
        "filters": {
            "grades": sorted(set(g[0] for g in all_grades if g[0])),
            "categories": sorted(set(c[0] for c in all_categories if c[0]))
        }
    })

@api_bp.route('/formulas/<int:id>', methods=['GET'])
def get_formula_detail(id):
    """获取单个公式详情"""
    formula = Formula.query.get(id)
    if not formula:
        return jsonify({"success": False, "message": "公式不存在"}), 404
    
    # 返回完整数据（包含 grade）
    data = formula.to_dict()
    data['grade'] = formula.grade
    return jsonify({"success": True, "data": data})

@api_bp.route('/formulas/search', methods=['POST'])
def search_formulas():
    """语义搜索公式"""
    from app.services.nlp_service import nlp_engine
    
    query_text = request.json.get('query', '').strip()
    category = request.json.get('category', '')
    grade = request.json.get('grade', '')

    if not query_text:
        return jsonify({"success": False, "message": "查询内容为空"})
        
    results = nlp_engine.search_formulas(query_text, category=category, grade=grade, top_k=5, threshold=0.5)
    
    if not results:
        return jsonify({"success": True, "data": [], "message": "无匹配公式"})
    
    return jsonify({"success": True, "data": results})

@api_bp.route('/formulas/explain', methods=['POST'])
def explain_formula():
    """公式智能助教：讲解 & 出题
    
    explain 模式：优先读取本地缓存，空则调 LLM 并回写数据库
    example 模式：每次重新生成
    """
    from app.services.llm_service import generate_formula_content
    from app.services.answer_engine import save_question_to_db

    data = request.json
    fid = data.get('id')
    mode = data.get('type', 'explain')

    if not fid:
        return jsonify({"success": False, "message": "Missing ID"})

    formula = Formula.query.get(fid)
    if not formula:
        return jsonify({"success": False, "message": "Formula not found"})

    # explain 模式：本地有缓存则同步返回（毫秒级）
    if mode == 'explain' and formula.explanation:
        return jsonify({"success": True, "data": formula.explanation, "source": "cache"})

    # 在 WSGI 线程中提取上下文（涉及 ORM），后台线程只做纯计算
    ctx = {
        "name": formula.name,
        "formula": formula.formula_text,
        "grade": formula.grade,
        "category": formula.category
    }
    formula_category = formula.category
    formula_id = formula.id
    app = current_app._get_current_object()

    def _explain():
        result = generate_formula_content(ctx, mode)
        if mode == 'example':
            if isinstance(result, dict):
                q_id = save_question_to_db(
                    question=result.get('question'),
                    answer=result.get('answer'),
                    reason=result.get('reason'),
                    options=result.get('options', []),
                    category=formula_category
                )
                return {"success": True, "data": result, "db_id": q_id}
            else:
                return {"success": False, "message": "生成格式错误"}
        else:
            # explain 模式：LLM 结果回写数据库
            if result:
                f = db.session.get(Formula, formula_id)
                if f:
                    f.explanation = result
                    db.session.commit()
            return {"success": True, "data": result}

    task_id = task_mgr.submit(_explain, app=app, owner=_get_owner())
    return jsonify({"success": True, "task_id": task_id}), 202

# === 单词消消乐 API ===
@api_bp.route('/words', methods=['GET'])
def get_random_words():
    """获取随机单词对"""
    try:
        count = request.args.get('count', 10, type=int)
        
        # 1. 查出最大 ID
        max_id = db.session.query(func.max(Vocabulary.id)).scalar()
        if not max_id:
            return jsonify({"success": True, "data": []})

        # 2. 生成随机 ID 列表
        import random
        random_ids = random.sample(range(1, max_id + 1), min(count, max_id))
        
        # 3. 使用 ID 列表精确查找
        words = Vocabulary.query.filter(Vocabulary.id.in_(random_ids)).all()
        
        # 如果查询结果数量不足，可以考虑补充一些数据，但这里简化处理

        return jsonify({
            "success": True,
            "data": [w.to_dict() for w in words]
        })
    except Exception as e:
        return jsonify({"success": False, "message": str(e)})

@api_bp.route('/words/search', methods=['GET'])
def search_vocabulary():
    """模糊搜索单词字典"""
    try:
        keyword = request.args.get('keyword', '').strip()
        if not keyword:
            return jsonify({"success": True, "data": []})
        
        words = Vocabulary.query.filter(
            or_(
                Vocabulary.word.like(f"%{keyword}%"),
                Vocabulary.definition.like(f"%{keyword}%")
            )
        ).limit(20).all()
        
        return jsonify({
            "success": True,
            "data": [w.to_dict() for w in words]
        })
    except Exception as e:
        return jsonify({"success": False, "message": str(e)})

# === 成语 PK API ===
@api_bp.route('/idioms/random', methods=['GET'])
def get_random_idioms():
    """获取随机成语"""
    try:
        count = request.args.get('count', 10, type=int)

        # 1. 查出最大 ID
        max_id = db.session.query(func.max(Idiom.id)).scalar()
        if not max_id:
            return jsonify({"success": True, "data": []})

        # 2. 生成随机 ID 列表
        import random
        random_ids = random.sample(range(1, max_id + 1), min(count, max_id))
        
        # 3. 使用 ID 列表精确查找
        idioms = Idiom.query.filter(Idiom.id.in_(random_ids)).all()
        
        return jsonify({
            "success": True,
            "data": [idiom.to_dict() for idiom in idioms]
        })
    except Exception as e:
        return jsonify({"success": False, "message": str(e)})

@api_bp.route('/idioms/all', methods=['GET'])
def get_all_idioms():
    """获取成语列表（支持分页）"""
    try:
        offset = request.args.get('offset', 0, type=int)
        limit = request.args.get('limit', 50, type=int)
        idioms = Idiom.query.order_by(Idiom.id).offset(offset).limit(limit).all()
        
        return jsonify({
            "success": True,
            "data": [idiom.to_dict() for idiom in idioms]
        })
    except Exception as e:
        return jsonify({"success": False, "message": str(e)})

@api_bp.route('/idioms/<int:id>', methods=['GET'])
def get_idiom_detail(id):
    """获取单个成语详情"""
    try:
        idiom = Idiom.query.get(id)
        if not idiom:
            return jsonify({"success": False, "message": "成语不存在"}), 404
            
        return jsonify({
            "success": True,
            "data": idiom.to_dict()
        })
    except Exception as e:
        return jsonify({"success": False, "message": str(e)})

@api_bp.route('/idioms/search', methods=['GET'])
def search_idioms():
    """搜索成语"""
    try:
        keyword = request.args.get('keyword', '').strip()
        if not keyword:
            return jsonify({"success": True, "data": []})
            
        # 移除关键词中的所有空格以便于拼音无缝搜索
        clean_keyword = keyword.replace(' ', '')
        
        idioms = Idiom.query.filter(
            or_(
                Idiom.word.like(f"%{keyword}%"),
                Idiom.abbreviation.like(f"{clean_keyword}%"),
                func.replace(Idiom.pinyin_r, ' ', '').like(f"{clean_keyword}%")
            )
        ).limit(20).all()
        
        return jsonify({
            "success": True,
            "data": [idiom.to_dict() for idiom in idioms]
        })
    except Exception as e:
        return jsonify({"success": False, "message": str(e)})

@api_bp.route('/idioms/validate_chain', methods=['GET'])
def validate_idiom_chain():
    """
    验证成语接龙是否合法 (同音不同字)
    参数:
    - word: 用户输入的成语
    - target: 需要匹配的起始拼音 (上一个成语最后一个字的读音，无声调)
    """
    try:
        word = request.args.get('word', '').strip()
        target_pinyin = request.args.get('target', '').strip()
        
        if not word or not target_pinyin:
            return jsonify({"success": False, "message": "缺少必要的验证参数"}), 400
            
        # 1. 在词库中查找该成语
        idiom = Idiom.query.filter_by(word=word).first()
        if not idiom:
            return jsonify({
                "success": True, 
                "valid": False, 
                "reason": f"词库中未收录成语「{word}」"
            })
            
        # 2. 直接获取数据库中已有的首尾拼音字段
        first_pinyin = idiom.first
        last_pinyin = idiom.last
        
        if not first_pinyin or not last_pinyin:
            return jsonify({
                 "success": True,
                 "valid": False,
                 "reason": f"成语「{word}」拼音数据缺失，无法判定"
            })
        
        # 3. 比对目标拼音
        if first_pinyin.lower() == target_pinyin.lower():
            data = idiom.to_dict()
            data['last_pinyin'] = last_pinyin # 下拉提示
            
            return jsonify({
                "success": True,
                "valid": True,
                "data": data
            })
        else:
            return jsonify({
                "success": True,
                "valid": False,
                "reason": f"「{word}」读音为 {first_pinyin}，不满足以 {target_pinyin} 开头的条件"
            })
            
    except Exception as e:
        return jsonify({"success": False, "message": str(e)})